from moviepy.editor import VideoFileClip, AudioFileClip
from Utils.ErrorDialog import ErrorDialog
from pydub import AudioSegment

from docx import Document
from docx.shared import Inches
from pdf2docx import Converter as convert

from fpdf import FPDF

from PIL import Image

import os

class Converter:
    def __init__(self):
        ...

    def convert_image(self, files, convert_to, additional_data=None, *args):
        print(convert_to)
        for file_path in files:
            try:
                original_image = Image.open(file_path)

                if additional_data:
                    upscale_factor = 1
                    if "upscale 2x" in additional_data:
                        upscale_factor = 2
                    elif "upscale 4x" in additional_data:
                        upscale_factor = 4

                    new_size = (original_image.width * upscale_factor, original_image.height * upscale_factor)
                    resized_image = original_image.resize(new_size, Image.BICUBIC)
                else:
                    resized_image = original_image

                file_name, file_ext = os.path.splitext(file_path)
                output_file_path = f"{file_name}_{convert_to}{file_ext}"

                resized_image.save(output_file_path + f'.{convert_to.lower()}', format=convert_to.lower())

            except Exception as e:
                ErrorDialog.show_error_dialog("Erro", f"Erro ao converter a imagem {file_path}: {e}", "tente novamente!")
                return False
        
        return True
    
    def convert_video(self, files, convert_to, additional_data=None, *args):
        codec_mapping = {
            'mp4': 'libx264',
            'avi': 'libxvid',
            'webm': 'libvpx',
            'ogv': 'libtheora',
            'mov': 'libx264'
        }

        audio_codec_mapping = {
            'mp4': 'aac',
            'avi': 'libmp3lame',
            'webm': 'libvorbis',
            'ogv': 'libvorbis',
            'mov': 'aac'
        }
        
        for file_path in files:
            try:
                video_clip = VideoFileClip(file_path)
                
                file_name, file_ext = os.path.splitext(file_path)
                output_file_path = f"{file_name}_converted.{convert_to.lower()}"
                
                codec = codec_mapping.get(convert_to.lower(), None)
                audio_codec = audio_codec_mapping.get(convert_to.lower(), 'aac')

                if codec is None:
                    raise ValueError(f"Formato de vídeo não suportado: {convert_to}")
                
                video_clip.write_videofile(
                    output_file_path,
                    codec=codec,
                    audio_codec=audio_codec,
                    temp_audiofile=f"{file_name}_temp_audio.m4a",
                    remove_temp=False,
                    verbose=True,
                    threads=4
                )
                
                if additional_data and additional_data.get('Render Audio Separately', False):
                    audio_output_path = f"{file_name}_audio.mp3"
                    video_clip.audio.write_audiofile(audio_output_path)

            except Exception as e:
                ErrorDialog.show_error_dialog("Erro", f"Erro ao converter o vídeo {file_path}: {e}", "Tente Novamente")
                return False
        
        return True

    def convert_audio(self, files, convert_to, additional_data=None, *args):
        format_mapping = {
            'mp3': 'mp3',
            'wav': 'wav',
            'ogg': 'ogg',
            'flac': 'flac',
            'aac': 'aac'
        }

        for file_path in files:
            try:
                audio = AudioSegment.from_file(file_path)
                
                file_name, file_ext = os.path.splitext(file_path)
                output_file_path = f"{file_name}_converted.{convert_to.lower()}"
                
                format_to = format_mapping.get(convert_to.lower(), None)
                if format_to is None:
                    raise ValueError(f"Formato de áudio não suportado: {convert_to}")
                audio.export(output_file_path, format=format_to)

            except Exception as e:
                ErrorDialog.show_error_dialog("Erro", f"Erro ao converter o áudio {file_path}: {e}", "Verifique se o programa FFMPEG está corretamente instalado")
                return False

        return True
    
    def convert_document(self, files, convert_to, additional_data, output_path):
        try:
            if not files:
                ErrorDialog.show_error_dialog("Erro", "Sem arquivos para converter.")
                return False

            if convert_to not in ['PDF', 'DOCX']:
                ErrorDialog.show_error_dialog("Erro", "Tipo de conversão não suportada.")
                return False

            output_file = os.path.join(output_path, f"{additional_data}.{convert_to.lower()}")

            if all(file.endswith(('.png', '.jpg', '.jpeg', '.webp')) for file in files):
                if convert_to == 'PDF':
                    pdf = FPDF()
                    for file in files:
                        pdf.add_page()
                        pdf.image(file, x=10, y=10, w=190)  # Ajuste conforme necessário
                    pdf.output(output_file)
                    return True
                
                elif convert_to == 'DOCX':
                    doc = Document()
                    for file in files:
                        doc.add_picture(file, width=Inches(2.0))  # Ajuste conforme necessário
                    doc.save(output_file)
                    return True

            elif all(file.endswith('.pdf') for file in files) and convert_to == 'DOCX':
                doc = Document()
                for pdf_file in files:
                    converter = convert(pdf_file=pdf_file)
                    converter.convert(output_file, start=0, end=None)
                    converter.close()
                    return True

            elif all(file.endswith('.docx') for file in files) and convert_to == 'PDF':
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)
                
                for docx_file in files:
                    doc = Document(docx_file)
                    
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            pdf.add_page()
                            pdf.multi_cell(0, 10, text)
                    
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            image = rel.target_part
                            image_bytes = image.blob
                            image_file = os.path.join(output_path, "temp_image.png")
                            
                            with open(image_file, "wb") as img_file:
                                img_file.write(image_bytes)
                            
                            pdf.add_page()
                            pdf.image(image_file, x=10, y=10, w=190)  # Ajuste conforme necessário
                            
                            os.remove(image_file)  # Remove o arquivo de imagem temporário
                
                pdf.output(output_file)
                return True

            else:
                ErrorDialog.show_error_dialog("Erro", "Arquivos não suportados ou erro de conversão.")
                return False

        except Exception as e:
            ErrorDialog.show_error_dialog("Erro de conversão", str(e))
            return False